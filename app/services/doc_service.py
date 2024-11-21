from docx import Document
from docx.shared import Pt  # 글자 크기
from docx.oxml.ns import qn

def replace_markers(template_path: str, output_path: str, data: dict):
    """
    템플릿 Word 문서에서 {{마커}}를 데이터로 치환하고, 마커가 없는 텍스트는 기본 스타일(맑은 고딕, 10pt)로 설정합니다.
    
    Args:
        template_path (str): 템플릿 파일 경로
        output_path (str): 결과 파일 경로
        data (dict): 치환할 데이터 (마커와 값)
    """
    # 템플릿 문서 열기
    doc = Document(template_path)

    # 문단 내 텍스트에서 마커를 치환
    for paragraph in doc.paragraphs:
        for key, value in data.items():
            marker = f"{{{{{key}}}}}"  # {{마커}} 형식
            if marker in paragraph.text:
                # 기존 텍스트를 마커 기준으로 분리
                parts = paragraph.text.split(marker)
                paragraph.text = ""  # 기존 텍스트 초기화

                for i, part in enumerate(parts):
                    # 마커 전후의 텍스트를 추가 (기본 스타일 적용)
                    if part:
                        run = paragraph.add_run(part)
                        run.font.name = "맑은 고딕"
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), "맑은 고딕")
                        run.font.size = Pt(10)  # 기본 글자 크기

                    # 마커에 해당하는 텍스트만 추가 및 스타일 적용
                    if i < len(parts) - 1:
                        run = paragraph.add_run(str(value))
                        run.font.name = "맑은 고딕"  # 폰트 설정
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), "맑은 고딕")  # 한글 폰트 설정

                        # 스타일 조건 적용
                        if key.endswith("_bold"):  # 굵은 스타일
                            run.bold = True
                            run.font.size = Pt(10)
                        elif key.endswith("_big_bold"):  # 큰 굵은 스타일
                            run.bold = True
                            run.font.size = Pt(16)
                        else:  # 기본 스타일
                            run.font.size = Pt(10)

    # 표의 셀에서도 마커 치환
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in data.items():
                    marker = f"{{{{{key}}}}}"
                    if marker in cell.text:
                        # 셀 텍스트를 마커 기준으로 분리
                        parts = cell.text.split(marker)
                        cell.text = ""  # 기존 텍스트 초기화

                        for i, part in enumerate(parts):
                            # 마커 전후의 텍스트 추가 (기본 스타일 적용)
                            if part:
                                run = cell.add_paragraph().add_run(part)
                                run.font.name = "맑은 고딕"
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), "맑은 고딕")
                                run.font.size = Pt(10)  # 기본 글자 크기

                            # 마커에 해당하는 텍스트만 추가 및 스타일 적용
                            if i < len(parts) - 1:
                                run = cell.add_paragraph().add_run(str(value))
                                run.font.name = "맑은 고딕"
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), "맑은 고딕")
                                
                                if key.endswith("_bold"):
                                    run.bold = True
                                    run.font.size = Pt(10)
                                elif key.endswith("_big_bold"):
                                    run.bold = True
                                    run.font.size = Pt(16)
                                else:
                                    run.font.size = Pt(10)

    # 결과 파일 저장
    doc.save(output_path)
